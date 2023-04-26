import requests
import datetime
from docx import Document

filename = '/home/lucky/Downloads/aptley/DISCReport/TemplateD.docx'
document = Document()

def replace_string(filename, oldText ,newText):
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if oldText in paragraph.text:
                        print('found and replacing')
                        paragraph.text = paragraph.text.replace(oldText, newText)
    return 1


def getArticle(topic):
  headers = {
    'Content-Type': 'application/json',
    'Authorization': 'Bearer changeme',
  }

  json_data = {
      'model': 'text-davinci-003',
      'prompt': "rephrase the following text to be more readable: " + topic,
      'temperature': .7,
      'max_tokens': 500,
  }
  response = requests.post('https://api.openai.com/v1/completions', headers=headers, json=json_data)
  article = response.json()
  print('request: ' + topic)
  print('reponse: ' + article['choices'][0]['text'].strip())
#   replace_string('/home/lucky/Downloads/aptley/DISCReport/TemplateC.docx', topic.strip(), article['choices'][0]['text'].strip())
  print('processing...')

a= [
'provide 3 muliplt choice question with answere for "python: Introduction"',
'provide 3 muliplt choice question with answere for "python: Installing Python & PyCharm"',
'provide 3 muliplt choice question with answere for "python: Setup & Hello World"',
'provide 3 muliplt choice question with answere for "python: Drawing a Shape"',
'provide 3 muliplt choice question with answere for "python: Variables & Data Types"',
'provide 3 muliplt choice question with answere for "python: Working With Strings"',
'provide 3 muliplt choice question with answere for "python: Working With Numbers"',
'provide 3 muliplt choice question with answere for "python: Getting Input From Users"',
'provide 3 muliplt choice question with answere for "python: Building a Basic Calculator"',
'provide 3 muliplt choice question with answere for "python: Mad Libs Game"',
'provide 3 muliplt choice question with answere for "python: Lists"',
'provide 3 muliplt choice question with answere for "python: List Functions"',
'provide 3 muliplt choice question with answere for "python: Tuples"',
'provide 3 muliplt choice question with answere for "python: Functions"',
'provide 3 muliplt choice question with answere for "python: Return Statement"',
'provide 3 muliplt choice question with answere for "python: If Statements"',
'provide 3 muliplt choice question with answere for "python: If Statements & Comparisons"',
'provide 3 muliplt choice question with answere for "python: Building a better Calculator"',
'provide 3 muliplt choice question with answere for "python: Dictionaries"',
'provide 3 muliplt choice question with answere for "python: While Loop"',
'provide 3 muliplt choice question with answere for "python: Building a Guessing Game"',
'provide 3 muliplt choice question with answere for "python: For Loops"',
'provide 3 muliplt choice question with answere for "python: Exponent Function"',
'provide 3 muliplt choice question with answere for "python: 2D Lists & Nested Loops"',
'provide 3 muliplt choice question with answere for "python: Building a Translator"',
'provide 3 muliplt choice question with answere for "python: Comments"',
'provide 3 muliplt choice question with answere for "python: Try / Except"',
'provide 3 muliplt choice question with answere for "python: Reading Files"',
'provide 3 muliplt choice question with answere for "python: Writing to Files"',
'provide 3 muliplt choice question with answere for "python: Modules & Pip"',
'provide 3 muliplt choice question with answere for "python: Classes & Objects"',
'provide 3 muliplt choice question with answere for "python: Building a Multiple Choice Quiz"',
'provide 3 muliplt choice question with answere for "python: Object Functions"',
'provide 3 muliplt choice question with answere for "python: Inheritance"',
'provide 3 muliplt choice question with answere for "python: Python Interpreter"'

    
]


print('total: ', len(a))

for i in a:
    getArticle(i)

